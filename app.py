from flask import Flask, request, render_template, send_file, redirect, url_for
import os
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
import tempfile
import google.generativeai as genai
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.linear_model import LogisticRegression
import pickle
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


app = Flask(__name__)

# Configure Gemini API
GEMINI_API_KEY = os.environ.get("AIzaSyDPgVkhl1jW_e89hc9opltcqHdpICYYv3g")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable not set.")
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")

feedback_data = []

# Machine Learning model initialization
if os.path.exists('feedback_model.pkl'):
    with open('feedback_model.pkl', 'rb') as model_file:
        feedback_model = pickle.load(model_file)
else:
    feedback_model = LogisticRegression()

def extract_text_from_pdf(file_path):
    """Try extracting text using PyMuPDF first, then fall back to OCR if needed."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text("text")
            text += page_text
        # If text is too short or empty, fallback to OCR
        if len(text.strip()) < 100:
            raise ValueError("Insufficient text from PyMuPDF; fallback to OCR.")
    except:
        # OCR fallback
        images = convert_from_path(file_path, dpi=300)
        for image in images:
            text += pytesseract.image_to_string(image)
    return text.strip()


def extract_text_from_image(image_file):
    from PIL import Image
    try:
        image = Image.open(image_file)
        return pytesseract.image_to_string(image)
    except Exception as e:
        return ""


import re

def extract_contact_info(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    top_section = "\n".join(lines[:30])  # Top ~30 lines, likely contains contact info
    full_text = "\n".join(lines)

    # Improved name extraction (first non-job-title line without digits or symbols)
    potential_names = [line for line in lines[:10] 
                       if not re.search(r'\d|@|http', line) and len(line.split()) <= 5]
    
    common_titles = ['curriculum', 'vitae', 'resume', 'developer', 'engineer', 'designer']
    name = ""
    for line in potential_names:
        if not any(word in line.lower() for word in common_titles):
            name = line.strip()
            break

    # Email
    emails = re.findall(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', top_section)
    emails = [e for e in emails if not any(dummy in e.lower() for dummy in ['noreply', 'help@', 'example'])]
    email = emails[0] if emails else ""

    # Phone (matches local, international, and common separators)
    phone_match = re.search(r'(\+?\d{1,3}[\s\-\.]?)?(\(?\d{2,4}\)?[\s\-\.]?)?\d{3,4}[\s\-\.]?\d{4}', top_section)
    phone = phone_match.group(0) if phone_match else ""

    # LinkedIn
    linkedin_match = re.search(r'https?://(www\.)?linkedin\.com/in/[^\s,<>")]+', full_text, re.I)
    linkedin = linkedin_match.group(0) if linkedin_match else ""

    # GitHub
    github_match = re.search(r'https?://(www\.)?github\.com/[^\s,<>")]+', full_text, re.I)
    github = github_match.group(0) if github_match else ""

    # Portfolio (any other URL not GitHub or LinkedIn)
    urls = re.findall(r'https?://[^\s,<>")]+', full_text)
    portfolio = ""
    for url in urls:
        if "github" not in url.lower() and "linkedin" not in url.lower():
            portfolio = url
            break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
    }


def analyze_similarity(resume_text, jd_text):
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, jd_text])
    sim_score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    return sim_score, vectorizer.get_feature_names_out()

def is_valid_resume(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    valid_extensions = ['.pdf', '.png', '.jpg', '.jpeg', '.bmp']
    
    if ext not in valid_extensions:
        return False

    try:
        if ext == '.pdf':
            resume_text = extract_text_from_pdf(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp']:
            resume_text = extract_text_from_image(file_path)
        else:
            return False
        
        if len(resume_text.strip()) < 100:
            return False

        # Check for typical resume keywords
        keywords = ['experience', 'education', 'skills', 'projects', 'summary']
        if not any(word in resume_text.lower() for word in keywords):
            return False

    except Exception:
        return False

    return True



def improve_resume(resume_text, jd):
    # Dynamically process the uploaded resume into structured sections
    sections = process_resume(resume_text)

    # Convert structured sections back into a clean plain-text format
    formatted_resume = ""
    for heading, lines in sections.items():
        formatted_resume += f"\n{heading.upper()}\n"
        for line in lines:
            formatted_resume += f"- {line.strip('*')}\n"

    # Gemini prompt using extracted structure
    prompt = f"""
You are an ATS resume optimization assistant.

Task:
Given the resume and job description below, generate a concise, professional, and ATS-friendly resume that fits strictly on ONE page (MAX ~450 words or ~50 bullet points), without adding any new information.

Guidelines:
- Use only the original resume content.
- Summarize and condense when needed.
- Keep section headings and subheadings as close to the original as possible.
- Optimize alignment to the job description.
- Output as plain text in resume format.

Resume (Structured Extract):
{formatted_resume}

Job Description:
{jd}

Return only the FINAL one-page resume (plain text, no HTML).
"""
    response = model.generate_content(prompt)
    return response.text.strip()[:3000]



def evaluate_resume(resume, jd):
    prompt = f"""
    Resume:
    {resume}

    Job Description:
    {jd}

    Identify and list:
    1. Matching skills
    2. Missing skills
    3. Strong points (highlight these)
    4. Weak points
    Return in a clean HTML format with <b> for highlights.
    """
    response = model.generate_content(prompt)
    return response.text
def process_resume(text):
    sections = {}
    current_section = None
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        # Detect headings: ALL CAPS and < 6 words
        if line.isupper() and len(line.split()) <= 6:
            current_section = line
            sections[current_section] = []
        elif current_section:
            # Detect subheading heuristically
            if re.match(r"^(?=.*[A-Za-z])([A-Z][A-Za-z0-9&,\-\. ]{3,})$", line):
                sections[current_section].append(f"**{line}**")  # Mark as subheading
            else:
                sections[current_section].append(line)
    return sections



def generate_docx(text, contact_info):
    doc = Document()

    name = contact_info.get("name", "")
    email = contact_info.get("email", "")
    phone = contact_info.get("phone", "")
    linkedin = contact_info.get("linkedin", "")
    github = contact_info.get("github", "")
    portfolio = contact_info.get("portfolio", "")

    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Name
    if name:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run(name)
        run.bold = True
        run.font.size = Pt(14)

    # Dynamically build contact line
    contact_parts = [part for part in [email, phone, linkedin, github, portfolio] if part]
    if contact_parts:
        contact_line = " | ".join(contact_parts)
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_para.add_run(contact_line)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(80, 80, 80)
        # Resume sections
        sections = process_resume(text)
        for section_title, lines in sections.items():
            title_para = doc.add_paragraph()
            run = title_para.add_run(section_title)
            run.bold = True
            run.font.size = Pt(10.5)
            title_para.paragraph_format.space_after = Pt(0)

# Add a horizontal line (bottom border)
            p = title_para._p  # Access the XML of the paragraph
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')        # Line thickness
            bottom.set(qn('w:space'), '1')     # Space between text and line
            bottom.set(qn('w:color'), 'auto')  # Use default color
            pBdr.append(bottom)
            pPr.append(pBdr)



            for line in lines:
                if line.startswith("**") and line.endswith("**"):
                    para = doc.add_paragraph(line.strip("**"))
                    para.paragraph_format.space_after = Pt(3)
                    para.paragraph_format.left_indent = Inches(0.05)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Keep subheading left-aligned
                    run = para.runs[0]
                    run.bold = True
                    run.font.size = Pt(9)
                else:
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.paragraph_format.space_after = Pt(3)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # <-- Justify normal text
                    run = para.runs[0]
                    run.font.size = Pt(9)


        

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)

    path = os.path.join(tempfile.gettempdir(), "final_resume.docx")
    doc.save(path)
    return path



def update_model(feedback_score, resume_text, jd_text):
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, jd_text])
    feedback_model.fit(vectors, [feedback_score])
    with open('feedback_model.pkl', 'wb') as model_file:
        pickle.dump(feedback_model, model_file)

@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('index.html', result=None)

@app.route('/evaluate', methods=['POST'])
def evaluate():
    resume_file = request.files['resume']
    jd = request.form['jd']
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        resume_file.save(tmp.name)
        resume_path = tmp.name

    if not is_valid_resume(resume_path):
        os.remove(resume_path)
        return "The uploaded file is not a valid resume. Please upload a real resume.", 400

    resume_text = extract_text_from_pdf(resume_path)
    result = evaluate_resume(resume_text, jd)

    os.remove(resume_path)
    return result


@app.route('/process', methods=['POST'])
def process():
    action = request.form['action']
    resume_file = request.files['resume']
    jd_text = request.form['job_description']

    ext = os.path.splitext(resume_file.filename)[1].lower()
    
    # Save the file temporarily to check its validity
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        resume_file.save(tmp.name)

        if not is_valid_resume(tmp.name):  # Validate the file
            return render_template('index.html', result="<h3 style='color:red;'>❌ Error: Invalid resume file. Please upload a valid resume (PDF or image with meaningful content).</h3>")

        # Proceed with text extraction and further processing if valid
        resume_text = ""
        try:
            if ext in ['.pdf']:
                resume_text = extract_text_from_pdf(tmp.name)
            elif ext in ['.png', '.jpg', '.jpeg', '.bmp']:
                resume_text = extract_text_from_image(tmp.name)
        except Exception as e:
            return render_template('index.html', result=f"<h3 style='color:red;'>❌ Failed to process the file. Error: {str(e)}</h3>")

    if not resume_text or len(resume_text.strip()) < 50:
        return render_template('index.html', result="<h3 style='color:red;'>❌ Error: Could not extract valid content from the uploaded resume. Please upload a clearer or properly formatted resume.</h3>")

    # Continue with the action based on 'evaluate', 'match', 'improve', etc.
    if action == 'evaluate':
        evaluation = evaluate_resume(resume_text, jd_text)
        return render_template('index.html', result=evaluation)

    elif action == 'match':
        sim_score, _ = analyze_similarity(resume_text, jd_text)
        return render_template('index.html', result=f"<h3>Match Score: <b>{sim_score:.2f}</b></h3>")

    elif action == 'improve':
        improved = improve_resume(resume_text, jd_text)
        if not improved or len(improved.strip()) < 100:
            return render_template('index.html', result="<h3 style='color:red;'>❌ Error: Failed to generate an improved resume. Please ensure the uploaded resume is clear and contains meaningful content.</h3>")
        contact_info = extract_contact_info(resume_text)
        if action == 'generate':
            improved_text = improve_resume(resume_text, jd_text)
            path = generate_docx(improved_text, contact_info)
            return send_file(path, as_attachment=True)

        docx_path = generate_docx(improved, contact_info)
        return send_file(docx_path, as_attachment=True, download_name="Improved_Resume.docx")

    # Handle feedback submission
    elif action == 'submit_feedback':
        feedback_score = int(request.form['feedback_score'])
        update_model(feedback_score, resume_text, jd_text)
        return redirect(url_for('index'))

    return redirect(url_for('index'))


if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=int(os.environ.get("PORT", 5000)))